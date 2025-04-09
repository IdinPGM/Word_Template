import docx
import pandas as pd

def fill_template(template_path, excel_path, output_path):
    """
    แก้ไขไฟล์ Word จากเทมเพลต โดยอิงข้อมูลจากไฟล์ Excel

    Args:
        template_path (str): เส้นทางไปยังไฟล์ Word เทมเพลต (.docx).
        excel_path (str): เส้นทางไปยังไฟล์ Excel ที่มีข้อมูล (.xlsx).
        output_path (str): เส้นทางสำหรับบันทึกไฟล์ Word ที่แก้ไขแล้ว (.docx).
    """
    try:
        # อ่านข้อมูลจากไฟล์ Excel
        df = pd.read_excel(excel_path)

        # ตรวจสอบว่ามีอย่างน้อยหนึ่งแถวข้อมูล
        if df.empty:
            print("ไม่พบข้อมูลในไฟล์ Excel.")
            return

        # โหลดไฟล์ Word เทมเพลต
        doc = docx.Document(template_path)

        # วนลูปผ่านแต่ละแถวใน DataFrame (แต่ละชุดข้อมูล)
        for index, row in df.iterrows():
            # สร้างสำเนาของเอกสารเทมเพลตสำหรับแต่ละชุดข้อมูล
            new_doc = docx.Document(template_path)

            # แทนที่ข้อความในเอกสาร
            for paragraph in new_doc.paragraphs:
                for key, value in row.items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

            for table in new_doc.tables:
                for row_table in table.rows:
                    for cell in row_table.cells:
                        for key, value in row.items():
                            placeholder = "{{" + key + "}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))

            # สร้างชื่อไฟล์เอาต์พุตใหม่ (เช่น output_row1.docx, output_row2.docx)
            output_filename = f"{output_path.replace('.docx', '')}_row{index + 1}.docx"
            new_doc.save(output_filename)
            print(f"สร้างไฟล์: {output_filename} สำเร็จ")

        print("การแก้ไขเทมเพลตเสร็จสมบูรณ์")

    except FileNotFoundError:
        print("ไม่พบไฟล์เทมเพลต หรือ ไฟล์ Excel โปรดตรวจสอบเส้นทางไฟล์อีกครั้ง.")
    except Exception as e:
        print(f"เกิดข้อผิดพลาด: {e}")

if __name__ == "__main__":
    template_file = "template.docx"  # เปลี่ยนเป็นชื่อไฟล์เทมเพลตของคุณ
    input_excel = "input.xlsx"      # เปลี่ยนเป็นชื่อไฟล์ Excel ของคุณ
    output_file_prefix = "output.docx" # คำนำหน้าชื่อไฟล์ output (จะมีการเพิ่ม _row[number].docx ต่อท้าย)

    fill_template(template_file, input_excel, output_file_prefix)